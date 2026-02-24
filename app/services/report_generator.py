"""
Report Generator - Generación de Reportes PDF y DOCX
SOC Agent - Fase 2

Genera reportes profesionales de investigaciones SOC con:
- Resumen ejecutivo
- IOCs analizados con resultados
- Gráficos de riesgo
- Técnicas MITRE detectadas
- Recomendaciones
- Timeline de la investigación
"""
import os
import json
import logging
from datetime import datetime
from typing import Dict, List, Optional, Any
from io import BytesIO

logger = logging.getLogger(__name__)


class ReportGenerator:
    """
    Generador de reportes profesionales para investigaciones SOC.

    Soporta:
    - PDF: Reporte visual con gráficos
    - DOCX: Reporte editable
    """

    # Colores corporativos SOC
    COLORS = {
        'primary': '#4F46E5',  # Indigo
        'secondary': '#6366F1',  # Indigo lighter
        'danger': '#DC2626',  # Red
        'warning': '#F59E0B',  # Amber
        'success': '#059669',  # Green
        'info': '#0284C7',  # Sky
        'dark': '#1F2937',  # Gray 800
        'light': '#F3F4F6',  # Gray 100
        'white': '#FFFFFF',
    }

    RISK_COLORS = {
        'CRÍTICO': '#DC2626',
        'ALTO': '#EA580C',
        'MEDIO': '#CA8A04',
        'BAJO': '#059669',
        'LIMPIO': '#2563EB',
    }

    def __init__(self):
        self.session_manager = None

    def _get_session_manager(self):
        if self.session_manager is None:
            from app.services.session_manager import SessionManager
            self.session_manager = SessionManager()
        return self.session_manager

    # =========================================================================
    # GENERACIÓN DE PDF
    # =========================================================================

    def generate_pdf(self, session_id: int, include_api_details: bool = False) -> Optional[BytesIO]:
        """
        Genera reporte PDF de una sesión de investigación.

        Args:
            session_id: ID de la sesión
            include_api_details: Si incluir datos crudos de APIs

        Returns:
            BytesIO con el PDF generado o None si hay error
        """
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.colors import HexColor
            from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
                PageBreak, Image, ListFlowable, ListItem
            )
            from reportlab.lib import colors

            # Obtener datos de la sesión
            sm = self._get_session_manager()
            session = sm.get_session(session_id)

            if not session:
                logger.error(f"Session {session_id} not found")
                return None

            session_data = sm.export_session_json(session_id)
            if not session_data:
                return None

            # Crear buffer para el PDF
            buffer = BytesIO()

            # Crear documento
            doc = SimpleDocTemplate(
                buffer,
                pagesize=letter,
                rightMargin=0.75 * inch,
                leftMargin=0.75 * inch,
                topMargin=0.75 * inch,
                bottomMargin=0.75 * inch
            )

            # Estilos
            styles = getSampleStyleSheet()

            # Estilos personalizados
            styles.add(ParagraphStyle(
                name='ReportTitle',
                parent=styles['Title'],
                fontSize=24,
                textColor=HexColor(self.COLORS['primary']),
                spaceAfter=20,
                alignment=TA_CENTER
            ))

            styles.add(ParagraphStyle(
                name='SectionHeader',
                parent=styles['Heading1'],
                fontSize=16,
                textColor=HexColor(self.COLORS['primary']),
                spaceBefore=20,
                spaceAfter=10,
                borderWidth=1,
                borderColor=HexColor(self.COLORS['primary']),
                borderPadding=5
            ))

            styles.add(ParagraphStyle(
                name='SubHeader',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=HexColor(self.COLORS['dark']),
                spaceBefore=15,
                spaceAfter=8
            ))

            styles.add(ParagraphStyle(
                name='BodyText',
                parent=styles['Normal'],
                fontSize=10,
                textColor=HexColor(self.COLORS['dark']),
                alignment=TA_JUSTIFY,
                spaceAfter=8
            ))

            styles.add(ParagraphStyle(
                name='SmallText',
                parent=styles['Normal'],
                fontSize=8,
                textColor=HexColor('#6B7280')
            ))

            # Contenido del reporte
            story = []

            # === PORTADA ===
            story.append(Spacer(1, 2 * inch))
            story.append(Paragraph("REPORTE DE INVESTIGACION SOC", styles['ReportTitle']))
            story.append(Spacer(1, 0.3 * inch))
            story.append(Paragraph(
                session_data['session'].get('title', 'Investigación de Seguridad'),
                styles['Heading2']
            ))
            story.append(Spacer(1, 0.5 * inch))

            # Info de la sesión
            session_info = session_data['session']
            risk_level = session_info.get('highest_risk_level', 'N/A')
            risk_color = self.RISK_COLORS.get(risk_level, self.COLORS['info'])

            # Tabla de metadata
            meta_data = [
                ['Fecha de Inicio', session_info.get('created_at', 'N/A')[:19].replace('T', ' ')],
                ['Estado', session_info.get('status', 'N/A').upper()],
                ['IOCs Analizados', str(session_info.get('total_iocs', 0))],
                ['Mensajes', str(session_info.get('total_messages', 0))],
                ['Nivel de Riesgo Máximo', risk_level],
            ]

            meta_table = Table(meta_data, colWidths=[2.5 * inch, 3.5 * inch])
            meta_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), HexColor(self.COLORS['light'])),
                ('TEXTCOLOR', (0, 0), (-1, -1), HexColor(self.COLORS['dark'])),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#E5E7EB')),
            ]))

            story.append(meta_table)
            story.append(Spacer(1, 0.5 * inch))

            # Fecha de generación
            story.append(Paragraph(
                f"Generado: {datetime.utcnow().strftime('%d %b %Y %H:%M UTC')}",
                styles['SmallText']
            ))

            story.append(PageBreak())

            # === RESUMEN EJECUTIVO ===
            story.append(Paragraph("RESUMEN EJECUTIVO", styles['SectionHeader']))

            # Generar resumen automático
            iocs = session_data.get('iocs', [])
            critical_count = sum(1 for i in iocs if i.get('risk_level') == 'CRÍTICO')
            high_count = sum(1 for i in iocs if i.get('risk_level') == 'ALTO')

            if critical_count > 0:
                summary = f"""
                Esta investigación ha identificado <b>{critical_count} indicador(es) CRÍTICO(S)</b> que requieren
                atención inmediata. Se analizaron un total de {len(iocs)} IOCs utilizando múltiples fuentes
                de threat intelligence.
                """
            elif high_count > 0:
                summary = f"""
                Se han detectado <b>{high_count} indicador(es) de ALTO riesgo</b> durante esta investigación.
                Se recomienda revisión prioritaria y posibles acciones de contención.
                """
            else:
                summary = f"""
                Se analizaron {len(iocs)} indicadores durante esta investigación. 
                No se detectaron amenazas críticas, aunque se recomienda monitoreo continuo.
                """

            story.append(Paragraph(summary, styles['BodyText']))
            story.append(Spacer(1, 0.3 * inch))

            # Estadísticas rápidas
            stats_data = [
                ['Métrica', 'Valor'],
                ['Total IOCs', str(len(iocs))],
                ['Críticos', str(critical_count)],
                ['Altos', str(high_count)],
                ['Fuentes consultadas',
                 str(len(set(s for i in iocs for s in (i.get('analysis_details', {}).get('sources_used') or []))))],
            ]

            stats_table = Table(stats_data, colWidths=[2 * inch, 1.5 * inch])
            stats_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), HexColor(self.COLORS['primary'])),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#E5E7EB')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, HexColor(self.COLORS['light'])]),
            ]))

            story.append(stats_table)
            story.append(Spacer(1, 0.5 * inch))

            # === IOCs ANALIZADOS ===
            story.append(Paragraph("INDICADORES DE COMPROMISO (IOCs)", styles['SectionHeader']))

            if iocs:
                for idx, ioc in enumerate(iocs, 1):
                    risk = ioc.get('risk_level', 'N/A')
                    risk_color = self.RISK_COLORS.get(risk, self.COLORS['info'])

                    # Header del IOC
                    story.append(Paragraph(
                        f"<b>IOC #{idx}:</b> {ioc.get('ioc_type', 'N/A').upper()}",
                        styles['SubHeader']
                    ))

                    # Tabla con datos del IOC
                    ioc_data = [
                        ['Valor', ioc.get('ioc_value', 'N/A')],
                        ['Tipo', ioc.get('ioc_type', 'N/A').upper()],
                        ['Riesgo', risk],
                        ['Score', f"{ioc.get('confidence_score', 'N/A')}/100"],
                        ['Rol', ioc.get('role', 'analyzed')],
                    ]

                    ioc_table = Table(ioc_data, colWidths=[1.5 * inch, 4.5 * inch])
                    ioc_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (0, -1), HexColor(self.COLORS['light'])),
                        ('TEXTCOLOR', (0, 0), (-1, -1), HexColor(self.COLORS['dark'])),
                        ('TEXTCOLOR', (1, 2), (1, 2), HexColor(risk_color)),  # Color del riesgo
                        ('FONTNAME', (1, 2), (1, 2), 'Helvetica-Bold'),
                        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#E5E7EB')),
                    ]))

                    story.append(ioc_table)

                    # Análisis LLM si existe
                    analysis_details = ioc.get('analysis_details', {})
                    llm_analysis = analysis_details.get('llm_analysis', {})

                    if llm_analysis and isinstance(llm_analysis, dict):
                        exec_summary = llm_analysis.get('executive_summary', '')
                        if exec_summary:
                            story.append(Spacer(1, 0.1 * inch))
                            story.append(Paragraph(f"<b>Análisis:</b> {exec_summary}", styles['BodyText']))

                        # Hallazgos clave
                        findings = llm_analysis.get('key_findings', [])
                        if findings:
                            story.append(Paragraph("<b>Hallazgos clave:</b>", styles['BodyText']))
                            for finding in findings[:5]:  # Máximo 5
                                story.append(Paragraph(f"• {finding}", styles['BodyText']))

                        # Recomendaciones
                        recommendations = llm_analysis.get('recommendations', [])
                        if recommendations:
                            story.append(Paragraph("<b>Recomendaciones:</b>", styles['BodyText']))
                            for rec in recommendations[:3]:  # Máximo 3
                                story.append(Paragraph(f"• {rec}", styles['BodyText']))

                    # Fuentes consultadas
                    sources = analysis_details.get('sources_used', [])
                    if sources:
                        story.append(Paragraph(
                            f"<b>Fuentes:</b> {', '.join(sources)}",
                            styles['SmallText']
                        ))

                    # MITRE ATT&CK
                    mitre = analysis_details.get('mitre_techniques', [])
                    if mitre:
                        story.append(Paragraph(
                            f"<b>MITRE ATT&CK:</b> {', '.join(mitre)}",
                            styles['SmallText']
                        ))

                    story.append(Spacer(1, 0.3 * inch))
            else:
                story.append(Paragraph("No se analizaron IOCs en esta sesión.", styles['BodyText']))

            # === TÉCNICAS MITRE ATT&CK ===
            all_mitre = set()
            for ioc in iocs:
                mitre = ioc.get('analysis_details', {}).get('mitre_techniques', [])
                if mitre:
                    all_mitre.update(mitre)

            if all_mitre:
                story.append(PageBreak())
                story.append(Paragraph("TECNICAS MITRE ATT&amp;CK DETECTADAS", styles['SectionHeader']))

                mitre_descriptions = {
                    'T1046': 'Network Service Discovery - Escaneo de servicios de red',
                    'T1071': 'Application Layer Protocol - Comunicación C2 via protocolos estándar',
                    'T1095': 'Non-Application Layer Protocol - Comunicación C2 via protocolos no estándar',
                    'T1595': 'Active Scanning - Escaneo activo de infraestructura',
                    'T1204': 'User Execution - Ejecución iniciada por usuario',
                    'T1059': 'Command and Scripting Interpreter - Ejecución de comandos/scripts',
                }

                for technique in sorted(all_mitre):
                    desc = mitre_descriptions.get(technique, 'Técnica detectada')
                    story.append(Paragraph(f"<b>{technique}</b>: {desc}", styles['BodyText']))

            # === RECOMENDACIONES GENERALES ===
            story.append(Spacer(1, 0.3 * inch))
            story.append(Paragraph("RECOMENDACIONES GENERALES", styles['SectionHeader']))

            recommendations = []
            if critical_count > 0:
                recommendations.extend([
                    "[!] URGENTE: Aislar sistemas potencialmente comprometidos",
                    "Bloquear IOCs críticos en firewall y proxy",
                    "Iniciar proceso de respuesta a incidentes",
                    "Revisar logs de los últimos 30 días para detectar actividad relacionada",
                ])
            elif high_count > 0:
                recommendations.extend([
                    "Monitorear tráfico hacia/desde IOCs de alto riesgo",
                    "Agregar IOCs a listas de vigilancia (watchlists)",
                    "Revisar sistemas que hayan contactado estos IOCs",
                ])
            else:
                recommendations.extend([
                    "Mantener monitoreo estándar",
                    "Actualizar firmas de detección",
                    "Continuar con procedimientos normales de seguridad",
                ])

            for rec in recommendations:
                story.append(Paragraph(f"• {rec}", styles['BodyText']))

            # === PIE DE PÁGINA / DISCLAIMER ===
            story.append(Spacer(1, 0.5 * inch))
            story.append(Paragraph(
                "Este reporte fue generado automáticamente por SOC Agent. "
                "Los datos provienen de múltiples fuentes de threat intelligence y análisis de IA. "
                "Se recomienda validación manual de hallazgos críticos.",
                styles['SmallText']
            ))

            # Construir PDF
            doc.build(story)

            buffer.seek(0)
            return buffer

        except Exception as e:
            logger.error(f"Error generating PDF: {e}")
            import traceback
            traceback.print_exc()
            return None

    # =========================================================================
    # GENERACIÓN DE DOCX
    # =========================================================================

    def generate_docx(self, session_id: int, include_api_details: bool = False) -> Optional[BytesIO]:
        """
        Genera reporte DOCX de una sesión de investigación.
        Usa python-docx (puro Python, sin dependencias externas).

        Args:
            session_id: ID de la sesión
            include_api_details: Si incluir datos crudos de APIs

        Returns:
            BytesIO con el DOCX generado o None si hay error
        """
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn

            # Obtener datos de la sesión
            sm = self._get_session_manager()
            session = sm.get_session(session_id)

            if not session:
                logger.error(f"Session {session_id} not found")
                return None

            session_data = sm.export_session_json(session_id)
            if not session_data:
                return None

            # Crear documento
            doc = Document()

            # Configurar estilos
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Arial'
            font.size = Pt(10)

            # Helper para colores
            def hex_to_rgb(hex_color):
                hex_color = hex_color.lstrip('#')
                return RGBColor(int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16))

            primary_color = hex_to_rgb(self.COLORS['primary'])

            # Helper para agregar tabla con shading
            def set_cell_shading(cell, color_hex):
                shading = cell._element.get_or_add_tcPr()
                shading_elem = shading.makeelement(qn('w:shd'), {
                    qn('w:fill'): color_hex.lstrip('#'),
                    qn('w:val'): 'clear'
                })
                shading.append(shading_elem)

            # ============================================================
            # PORTADA
            # ============================================================

            # Espaciado superior
            for _ in range(4):
                doc.add_paragraph()

            # Título
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run('REPORTE DE INVESTIGACION SOC')
            run.font.size = Pt(24)
            run.font.color.rgb = primary_color
            run.bold = True

            # Subtítulo
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subtitle.add_run(session_data['session'].get('title', 'Investigacion de Seguridad'))
            run.font.size = Pt(14)
            run.font.color.rgb = hex_to_rgb(self.COLORS['dark'])

            doc.add_paragraph()

            # Tabla de metadata
            session_info = session_data['session']
            risk_level = session_info.get('highest_risk_level', 'N/A')

            meta_rows = [
                ('Fecha de Inicio', (session_info.get('created_at', 'N/A') or 'N/A')[:19].replace('T', ' ')),
                ('Estado', (session_info.get('status', 'N/A') or 'N/A').upper()),
                ('IOCs Analizados', str(session_info.get('total_iocs', 0))),
                ('Mensajes', str(session_info.get('total_messages', 0))),
                ('Nivel de Riesgo Maximo', risk_level),
            ]

            table = doc.add_table(rows=len(meta_rows), cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            for i, (label, value) in enumerate(meta_rows):
                cell_label = table.rows[i].cells[0]
                cell_value = table.rows[i].cells[1]

                cell_label.text = label
                cell_label.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell_label, self.COLORS['light'])

                cell_value.text = value
                if label == 'Nivel de Riesgo Maximo' and risk_level in self.RISK_COLORS:
                    cell_value.paragraphs[0].runs[0].font.color.rgb = hex_to_rgb(self.RISK_COLORS[risk_level])
                    cell_value.paragraphs[0].runs[0].bold = True

            doc.add_paragraph()

            # Fecha de generación
            gen_date = doc.add_paragraph()
            gen_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = gen_date.add_run(f"Generado: {datetime.utcnow().strftime('%d %b %Y %H:%M UTC')}")
            run.font.size = Pt(8)
            run.font.color.rgb = hex_to_rgb('#6B7280')

            doc.add_page_break()

            # ============================================================
            # RESUMEN EJECUTIVO
            # ============================================================

            heading = doc.add_heading('RESUMEN EJECUTIVO', level=1)
            for run in heading.runs:
                run.font.color.rgb = primary_color

            iocs = session_data.get('iocs', [])
            critical_count = sum(
                1 for i in iocs if i.get('risk_level') == 'CRITICO' or i.get('risk_level') == 'CRÍTICO')
            high_count = sum(1 for i in iocs if i.get('risk_level') == 'ALTO')

            if critical_count > 0:
                summary_text = (
                    f"Esta investigacion ha identificado {critical_count} indicador(es) CRITICO(S) "
                    f"que requieren atencion inmediata. Se analizaron un total de {len(iocs)} IOCs "
                    f"utilizando multiples fuentes de threat intelligence."
                )
            elif high_count > 0:
                summary_text = (
                    f"Se han detectado {high_count} indicador(es) de ALTO riesgo durante esta investigacion. "
                    f"Se recomienda revision prioritaria y posibles acciones de contencion."
                )
            else:
                summary_text = (
                    f"Se analizaron {len(iocs)} indicadores durante esta investigacion. "
                    f"No se detectaron amenazas criticas, aunque se recomienda monitoreo continuo."
                )

            doc.add_paragraph(summary_text)

            # Tabla de estadísticas
            stats_table = doc.add_table(rows=5, cols=2)
            stats_data = [
                ('Total IOCs', str(len(iocs))),
                ('Criticos', str(critical_count)),
                ('Altos', str(high_count)),
                ('Fuentes consultadas', str(len(set(
                    s for i in iocs
                    for s in (i.get('analysis_details', {}).get('sources_used') or [])
                )))),
                ('Periodo', (session_info.get('created_at', 'N/A') or 'N/A')[:10]),
            ]

            for i, (label, value) in enumerate(stats_data):
                stats_table.rows[i].cells[0].text = label
                stats_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
                set_cell_shading(stats_table.rows[i].cells[0], self.COLORS['light'])
                stats_table.rows[i].cells[1].text = value

            doc.add_paragraph()

            # ============================================================
            # IOCs ANALIZADOS
            # ============================================================

            heading = doc.add_heading('INDICADORES DE COMPROMISO (IOCs)', level=1)
            for run in heading.runs:
                run.font.color.rgb = primary_color

            if iocs:
                for idx, ioc in enumerate(iocs, 1):
                    risk = ioc.get('risk_level', 'N/A')
                    risk_color_hex = self.RISK_COLORS.get(risk, self.COLORS['info'])

                    # Sub-header del IOC
                    doc.add_heading(
                        f"IOC #{idx}: {(ioc.get('ioc_type') or 'N/A').upper()}",
                        level=2
                    )

                    # Tabla con datos del IOC
                    ioc_table = doc.add_table(rows=4, cols=2)
                    ioc_rows = [
                        ('Valor', ioc.get('ioc_value', 'N/A')),
                        ('Tipo', (ioc.get('ioc_type') or 'N/A').upper()),
                        ('Riesgo', risk),
                        ('Score', f"{ioc.get('confidence_score', 'N/A')}/100"),
                    ]

                    for i, (label, value) in enumerate(ioc_rows):
                        cell_l = ioc_table.rows[i].cells[0]
                        cell_v = ioc_table.rows[i].cells[1]
                        cell_l.text = label
                        cell_l.paragraphs[0].runs[0].bold = True
                        set_cell_shading(cell_l, self.COLORS['light'])
                        cell_v.text = str(value)

                        # Color del riesgo
                        if label == 'Riesgo' and risk in self.RISK_COLORS:
                            cell_v.paragraphs[0].runs[0].font.color.rgb = hex_to_rgb(self.RISK_COLORS[risk])
                            cell_v.paragraphs[0].runs[0].bold = True

                    # Análisis LLM
                    analysis_details = ioc.get('analysis_details', {})
                    llm_analysis = analysis_details.get('llm_analysis', {})

                    if llm_analysis and isinstance(llm_analysis, dict):
                        exec_summary = llm_analysis.get('executive_summary', '')
                        if exec_summary:
                            p = doc.add_paragraph()
                            p.add_run('Analisis: ').bold = True
                            p.add_run(exec_summary)

                        findings = llm_analysis.get('key_findings', [])
                        if findings:
                            p = doc.add_paragraph()
                            p.add_run('Hallazgos clave:').bold = True
                            for finding in findings[:5]:
                                doc.add_paragraph(finding, style='List Bullet')

                        recommendations = llm_analysis.get('recommendations', [])
                        if recommendations:
                            p = doc.add_paragraph()
                            p.add_run('Recomendaciones:').bold = True
                            for rec in recommendations[:3]:
                                doc.add_paragraph(rec, style='List Bullet')

                    # Fuentes
                    sources = analysis_details.get('sources_used', [])
                    if sources:
                        p = doc.add_paragraph()
                        run = p.add_run(f"Fuentes: {', '.join(sources)}")
                        run.font.size = Pt(8)
                        run.font.color.rgb = hex_to_rgb('#6B7280')

                    # MITRE
                    mitre = analysis_details.get('mitre_techniques', [])
                    if mitre:
                        p = doc.add_paragraph()
                        run = p.add_run(f"MITRE ATT&CK: {', '.join(mitre)}")
                        run.font.size = Pt(8)
                        run.font.color.rgb = hex_to_rgb('#6B7280')

                    # Datos técnicos de APIs (opcional)
                    if include_api_details:
                        for api_name in ['virustotal_data', 'abuseipdb_data', 'greynoise_data', 'threatfox_data']:
                            api_data = analysis_details.get(api_name)
                            if api_data and isinstance(api_data, dict) and not api_data.get('error'):
                                p = doc.add_paragraph()
                                run = p.add_run(f"{api_name.replace('_data', '').upper()}: ")
                                run.bold = True
                                run.font.size = Pt(8)
                                detail_run = p.add_run(json.dumps(api_data, indent=2, default=str)[:500])
                                detail_run.font.size = Pt(7)

                    doc.add_paragraph()  # Espaciado
            else:
                doc.add_paragraph("No se analizaron IOCs en esta sesion.")

            # ============================================================
            # TÉCNICAS MITRE ATT&CK
            # ============================================================

            all_mitre = set()
            for ioc in iocs:
                mitre = ioc.get('analysis_details', {}).get('mitre_techniques', [])
                if mitre:
                    all_mitre.update(mitre)

            if all_mitre:
                doc.add_page_break()
                heading = doc.add_heading('TECNICAS MITRE ATT&CK DETECTADAS', level=1)
                for run in heading.runs:
                    run.font.color.rgb = primary_color

                mitre_descriptions = {
                    'T1046': 'Network Service Discovery',
                    'T1071': 'Application Layer Protocol - C2',
                    'T1095': 'Non-Application Layer Protocol - C2',
                    'T1595': 'Active Scanning',
                    'T1204': 'User Execution',
                    'T1059': 'Command and Scripting Interpreter',
                }

                for technique in sorted(all_mitre):
                    desc = mitre_descriptions.get(technique, 'Tecnica detectada')
                    p = doc.add_paragraph()
                    p.add_run(f"{technique}: ").bold = True
                    p.add_run(desc)

            # ============================================================
            # RECOMENDACIONES GENERALES
            # ============================================================

            heading = doc.add_heading('RECOMENDACIONES GENERALES', level=1)
            for run in heading.runs:
                run.font.color.rgb = primary_color

            if critical_count > 0:
                recs = [
                    "URGENTE: Aislar sistemas potencialmente comprometidos",
                    "Bloquear IOCs criticos en firewall y proxy",
                    "Iniciar proceso de respuesta a incidentes",
                    "Revisar logs de los ultimos 30 dias para detectar actividad relacionada",
                ]
            elif high_count > 0:
                recs = [
                    "Monitorear trafico hacia/desde IOCs de alto riesgo",
                    "Agregar IOCs a listas de vigilancia (watchlists)",
                    "Revisar sistemas que hayan contactado estos IOCs",
                ]
            else:
                recs = [
                    "Mantener monitoreo estandar",
                    "Actualizar firmas de deteccion",
                    "Continuar con procedimientos normales de seguridad",
                ]

            for rec in recs:
                doc.add_paragraph(rec, style='List Bullet')

            # ============================================================
            # FOOTER / DISCLAIMER
            # ============================================================

            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(
                "Este reporte fue generado automaticamente por SOC Agent. "
                "Los datos provienen de multiples fuentes de threat intelligence y analisis de IA. "
                "Se recomienda validacion manual de hallazgos criticos."
            )
            run.font.size = Pt(8)
            run.font.color.rgb = hex_to_rgb('#6B7280')

            p = doc.add_paragraph()
            run = p.add_run(f"Generado por SOC Agent - {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
            run.font.size = Pt(8)
            run.font.color.rgb = hex_to_rgb('#6B7280')

            # Guardar en buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"Error generating DOCX: {e}")
            import traceback
            traceback.print_exc()
            return None


# Instancia global
report_generator = ReportGenerator()